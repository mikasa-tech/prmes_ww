#!/usr/bin/env python3

try:
    import docx
    from docx import Document
    
    def extract_docx_content():
        """Extract content from the reference DOCX file"""
        doc_path = r"C:\Users\asus\Downloads\2022 SCHEME MAJOR PHASE-I PROJECT REVIEW-I (2).docx"
        
        try:
            doc = Document(doc_path)
            
            print("=" * 60)
            print("DOCUMENT STRUCTURE AND CONTENT")
            print("=" * 60)
            
            # Extract paragraphs
            print("\n--- PARAGRAPHS ---")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    print(f"Para {i+1}: {para.text}")
                    # Check formatting
                    for run in para.runs:
                        if run.bold:
                            print(f"  [BOLD]: {run.text}")
                        if run.font.size:
                            print(f"  [Font Size]: {run.font.size}")
            
            # Extract tables
            print("\n--- TABLES ---")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    print(f"  Row {row_idx+1}: {' | '.join(row_text)}")
            
            # Try to extract images/shapes info
            print("\n--- DOCUMENT SECTIONS ---")
            for section in doc.sections:
                print(f"Section margins: {section.top_margin}, {section.bottom_margin}")
                print(f"Page size: {section.page_width} x {section.page_height}")
            
        except FileNotFoundError:
            print(f"File not found: {doc_path}")
            print("Available DOCX files:")
            import os
            downloads = r"C:\Users\asus\Downloads"
            for file in os.listdir(downloads):
                if file.endswith('.docx'):
                    print(f"  - {file}")
                    
        except Exception as e:
            print(f"Error reading document: {e}")
            import traceback
            traceback.print_exc()
    
    if __name__ == "__main__":
        extract_docx_content()
        
except ImportError:
    print("python-docx not installed. Install with: pip install python-docx")